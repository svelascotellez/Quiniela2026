import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Styles
title_style = doc.styles['Title']
title_font = title_style.font
title_font.name = 'Arial'
title_font.size = Pt(22)
title_font.color.rgb = RGBColor(0x0f, 0x4c, 0x43) # Green from project
title_font.bold = True

h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Arial'
h1_font.size = Pt(16)
h1_font.color.rgb = RGBColor(0x0f, 0x4c, 0x43)
h1_font.bold = True

h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Arial'
h2_font.size = Pt(14)
h2_font.color.rgb = RGBColor(0x33, 0x33, 0x33)
h2_font.bold = True

# Content
doc.add_heading('🏆 LA QUINIELA DEFINITIVA - MUNDIAL NORTEAMÉRICA 2026 🌎⚽', 0)

p = doc.add_paragraph()
p.add_run('¡Prepárate para vivir el Mundial 2026 como nunca antes! Les presento ').bold = False
p.add_run('el sistema de quiniela más avanzado, emocionante y automatizado').bold = True
p.add_run(', diseñado para llevar la pasión del fútbol al siguiente nivel. Olvídate de llevar cuentas a mano o en libretas aburridas; la inteligencia de datos se encargará de todo.').bold = False

doc.add_heading('🌟 ¿Qué hace a esta Quiniela única en su clase?', level=1)

doc.add_heading('1. 🖥️ Plataforma Web Exclusiva y Siempre Disponible', level=2)
doc.add_paragraph('No más correos perdidos o archivos de WhatsApp que no encuentras. Contamos con nuestra propia aplicación web privada disponible 24/7 en la nube. Podrás descargar la plantilla oficial para participar y estar al tanto de los resultados.', style='List Bullet')

doc.add_heading('2. 📊 Plantilla Excel Premium y Súper Completa', level=2)
doc.add_paragraph('No rellenarás un archivo genérico, es un documento diseñado profesionalmente que incluye:')
doc.add_paragraph('Calendario Real y Actualizado: Fechas, estadios y horarios de los 104 partidos convertidos automáticamente a la hora de la Ciudad de México (CDMX).', style='List Bullet')
doc.add_paragraph('Fase de Grupos y Eliminatorias: Toda la estructura del torneo.', style='List Bullet')
doc.add_paragraph('Tablas de Posiciones Automáticas: El archivo hace los cálculos matemáticos para definir posiciones de grupo según se llenan los resultados de forma dinámica.', style='List Bullet')
doc.add_paragraph('Ranking FIFA Oficial: Una pestaña dedicada con el Ranking FIFA de todas las selecciones para ayudarte a tomar las mejores decisiones como experto.', style='List Bullet')

doc.add_heading('3. 🎯 Sistema de Puntos Estilo "March Madness"', level=2)
doc.add_paragraph('¡La emoción dura todo el torneo! No solo ganas puntos por atinarle a los marcadores exactos, sino que usamos un sistema mucho más entretenido:')
doc.add_paragraph('Ganas puntos si adivinas qué selecciones logran superar las barreras y clasificar a Dieciseisavos, Octavos, Cuartos, Semis y la Final.', style='List Bullet')
doc.add_paragraph('Podio de Honor: Un "bono gordo" de puntos si logras adivinar, desde antes que arranque el mundial, al Campeón, Subcampeón y Tercer Lugar exactos.', style='List Bullet')

doc.add_heading('4. 🥇 Clasificación (Leaderboard) Automática y Espectacular', level=2)
doc.add_paragraph('La joya de la corona. Con un solo clic de los administradores del sistema, un algoritmo evalúa todas nuestras predicciones contra la realidad y genera automáticamente:')
doc.add_paragraph('Un archivo de Resultados Espectacular: Con diseño verde esmeralda, que otorga automáticamente medallas de Oro 🏆, Plata 🥈 y Bronce 🥉 para el Top 3 de los participantes.', style='List Bullet')
doc.add_paragraph('Pestañas Individuales de Auditoría: El sistema genera a cada participante su propia pestaña detallando sus porcentajes de efectividad, dónde acertó y dónde falló miserablemente.', style='List Bullet')

doc.add_heading('5. 📱 ¡Reportes Inmediatos por WhatsApp!', level=2)
doc.add_paragraph('Para calentar los ánimos en nuestro grupo de chat, el sistema escupe un resumen inmediato y detallado del Top de posiciones para enterarnos quién va ganando la bolsa al momento.', style='List Bullet')

doc.add_heading('💡 ¿Por qué unirte a esta Quiniela?', level=1)
doc.add_paragraph('Cero Problemas de "Dedazo": Un programa de computadora inteligente audita todas las plantillas, por lo que no hay cálculos tramposos o humanos.', style='List Number')
doc.add_paragraph('Transparencia Total: Las reglas de puntos son claras y los aciertos están comprobables.', style='List Number')
doc.add_paragraph('Experiencia Premium: Es como participar en una liga deportiva profesional, no en una quiniela del montón.', style='List Number')

p = doc.add_paragraph()
p.add_run('¿Le entras al reto del 2026? 🚀🔥').bold = True
p.add_run(' ¡Pide la plantilla hoy mismo y demuestra qué tanto sabes de fútbol!')

doc.save('Presentacion_Quiniela_2026.docx')
print("Document generated!")
